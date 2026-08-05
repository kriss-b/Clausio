"""Extraction textuelle des candidatures.

Gère PDF, Word, Excel, texte et archives ZIP imbriquées. Chaque bloc extrait
conserve sa provenance (document, page, section) pour tracer les preuves.
Les PDF scannés (sans couche texte) sont signalés : l'OCR se ferait via Albert.
Dégradation gracieuse si une bibliothèque n'est pas installée.
"""
from __future__ import annotations

import zipfile
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class Bloc:
    document_nom: str
    texte: str
    page: int | None = None
    section: str = ""


@dataclass
class ResultatExtraction:
    blocs: list[Bloc] = field(default_factory=list)
    avertissements: list[str] = field(default_factory=list)


def extraire(chemin: Path, nom_affiche: str | None = None) -> ResultatExtraction:
    nom = nom_affiche or chemin.name
    suffixe = chemin.suffix.lower()
    res = ResultatExtraction()

    try:
        if suffixe == ".pdf":
            _pdf(chemin, nom, res)
        elif suffixe == ".docx":
            _docx(chemin, nom, res)
        elif suffixe == ".xlsx":
            _xlsx(chemin, nom, res)
        elif suffixe in {".txt", ".md", ".csv"}:
            res.blocs.append(Bloc(nom, chemin.read_text(encoding="utf-8", errors="replace")))
        elif suffixe == ".zip":
            _zip(chemin, nom, res)
        else:
            res.avertissements.append(f"{nom} : format {suffixe} non pris en charge.")
    except Exception as exc:  # noqa: BLE001
        res.avertissements.append(f"{nom} : échec d'extraction ({exc}).")
    return res


def _pdf(chemin: Path, nom: str, res: ResultatExtraction) -> None:
    try:
        import fitz  # pymupdf
    except ImportError:
        res.avertissements.append(f"{nom} : pymupdf absent, PDF non extrait.")
        return
    doc = fitz.open(chemin)
    vide = 0
    for i, page in enumerate(doc, start=1):
        txt = page.get_text().strip()
        if txt:
            res.blocs.append(Bloc(nom, txt, page=i))
        else:
            vide += 1
    if vide:
        res.avertissements.append(
            f"{nom} : {vide} page(s) sans texte — probablement scannées, OCR (Albert) requis."
        )


def _docx(chemin: Path, nom: str, res: ResultatExtraction) -> None:
    try:
        import docx  # python-docx
    except ImportError:
        res.avertissements.append(f"{nom} : python-docx absent, Word non extrait.")
        return
    d = docx.Document(str(chemin))
    paras = [p.text for p in d.paragraphs if p.text.strip()]
    if paras:
        res.blocs.append(Bloc(nom, "\n".join(paras)))
    for t_idx, table in enumerate(d.tables, start=1):
        lignes = [" | ".join(c.text for c in row.cells) for row in table.rows]
        if lignes:
            res.blocs.append(Bloc(nom, "\n".join(lignes), section=f"tableau {t_idx}"))


def _xlsx(chemin: Path, nom: str, res: ResultatExtraction) -> None:
    try:
        import openpyxl
    except ImportError:
        res.avertissements.append(f"{nom} : openpyxl absent, Excel non extrait.")
        return
    wb = openpyxl.load_workbook(chemin, read_only=True, data_only=True)
    for ws in wb.worksheets:
        lignes = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                lignes.append(" | ".join(cells))
        if lignes:
            res.blocs.append(Bloc(nom, "\n".join(lignes), section=f"feuille {ws.title}"))


def _zip(chemin: Path, nom: str, res: ResultatExtraction, profondeur: int = 0) -> None:
    if profondeur > 5:
        res.avertissements.append(f"{nom} : imbrication ZIP trop profonde, arrêt.")
        return
    import tempfile
    with zipfile.ZipFile(chemin) as zf, tempfile.TemporaryDirectory() as tmp:
        for membre in zf.namelist():
            if membre.endswith("/"):
                continue
            cible = Path(zf.extract(membre, tmp))
            sous = extraire(cible, nom_affiche=f"{nom}::{membre}")
            res.blocs.extend(sous.blocs)
            res.avertissements.extend(sous.avertissements)
